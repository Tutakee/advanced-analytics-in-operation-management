"""
merge_sections.py  —  Replace Sections 4 & 5 in Combined_Paper.docx with the
verified experimental draft, and append new APA references.

Output: papers/Combined_Paper_Revised.docx
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ── table border XML (copied from existing tables in Combined_Paper.docx) ──
_TBLPR = """<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:tblW w:w="9737" w:type="dxa"/>
  <w:tblBorders>
    <w:top    w:val="single" w:color="000000" w:sz="12"/>
    <w:left   w:val="single" w:color="000000" w:sz="12"/>
    <w:bottom w:val="single" w:color="000000" w:sz="12"/>
    <w:right  w:val="single" w:color="000000" w:sz="12"/>
    <w:insideH w:val="single" w:color="000000" w:sz="12"/>
    <w:insideV w:val="single" w:color="000000" w:sz="12"/>
  </w:tblBorders>
  <w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0"
             w:firstColumn="1" w:lastColumn="0"
             w:noHBand="0" w:noVBand="1"/>
</w:tblPr>"""


def _apply_borders(tbl):
    new_pr = etree.fromstring(_TBLPR)
    old_pr = tbl._tbl.find(qn("w:tblPr"))
    if old_pr is not None:
        tbl._tbl.remove(old_pr)
    tbl._tbl.insert(0, new_pr)


# ── paragraph factory (creates element but does NOT attach it to doc body) ─

def _make_para(doc, text, style="Normal", bold_prefix=None):
    """
    Create a paragraph element detached from the body.
    If bold_prefix is given, the first run is bold, the rest (text) is normal.
    """
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    elem = p._element
    doc.element.body.remove(elem)   # detach — we place it manually
    return elem


def _make_heading(doc, text, level):
    styles = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(text, style=styles[level])
    elem = p._element
    doc.element.body.remove(elem)
    return elem


def _make_table(doc, headers, rows_data):
    """Create a bordered table element detached from the body."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    _apply_borders(tbl)
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows_data):
        for j, v in enumerate(row):
            tbl.rows[i + 1].cells[j].text = v
    elem = tbl._tbl
    doc.element.body.remove(elem)
    return elem


def _make_caption(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    if p.runs:
        p.runs[0].bold = True
    elem = p._element
    doc.element.body.remove(elem)
    return elem


def _make_spacer(doc):
    p = doc.add_paragraph("", style="Normal")
    elem = p._element
    doc.element.body.remove(elem)
    return elem


# ── build section content as list of XML elements ─────────────────────────

def build_section_4(doc):
    D = doc
    elems = []
    A = elems.append

    A(_make_heading(D, "4. Experiments", 1))
    A(_make_para(D,
        "We evaluate the multi-group Wasserstein distributionally robust fair classifier "
        "(HDRFC-K) of Theorem 3.5 on a real-world insurance pricing dataset, in which the "
        "sensitive attribute is driver age partitioned into K = 5 groups. The experimental "
        "design is intentionally conservative: a single representative dataset, a transparent "
        "linear-program implementation, and a head-to-head comparison against the unconstrained "
        "Wasserstein support-vector machine recovered from the same LP by setting ρ = 0 and "
        "ζ = ∞. Our objective is to test whether the multi-group equal-opportunity constraints "
        "(8)–(10) materially reduce true positive rate (TPR) disparity at acceptable cost to "
        "predictive accuracy."
    ))

    A(_make_heading(D, "4.1 Dataset and Preprocessing", 2))
    A(_make_para(D,
        "We use the freMTPL2freq dataset (French Motor Third-Party Liability claim frequency, "
        "OpenML ID 41214; Duval & Charpentier, 2020), a public benchmark in actuarial science "
        "containing 678,013 one-year insurance policies. The binary outcome is y_i = +1 if at "
        "least one claim was filed during the policy year and y_i = −1 otherwise. Because the "
        "LP in Theorem 3.5 has O(N + Σ_{(k,l)} |I_k⁺| + |I_l⁺|) variables and constraints, "
        "we draw a stratified subsample of N = 20,000 policies, stratifying on the joint "
        "distribution of Y and the sensitive attribute A to preserve the empirical "
        "class-by-group composition. The subsample is split 80%/20% into a training set "
        "(N_tr = 16,000) and a held-out test set (N_te = 4,000); the positive rate is 5.0% "
        "(1,004 positives out of 20,000)."
    ))
    A(_make_para(D,
        "The sensitive attribute A_i ∈ {1,…,5} is constructed by binning driver age into the "
        "intervals [18, 25], [26, 40], [41, 55], [56, 70], and [71, ∞). Driver age itself is "
        "excluded from the feature vector x_i at training and test time, so that the classifier "
        "cannot trivially condition on the protected variable. After one-hot encoding the four "
        "categorical regressors (vehicle brand, vehicle gas, region, area code) and standardising "
        "the five numeric regressors (vehicle power, vehicle age, bonus-malus, density, exposure), "
        "the feature dimension is d = 42. Table 1 summarises the per-group sample sizes and base "
        "rates on the training fold."
    ))

    A(_make_caption(D, "Table 1. Age-group composition of the training fold (N_tr = 16,000)."))
    A(_make_table(D,
        ["Group k", "Age range", "N_k", "|I_k+|", "Positive rate"],
        [
            ["1", "18–25", "918",   "62",  "6.7%"],
            ["2", "26–40", "5,562", "238", "4.3%"],
            ["3", "41–55", "5,799", "305", "5.3%"],
            ["4", "56–70", "2,792", "141", "5.0%"],
            ["5", "71+",        "929",   "58",  "6.3%"],
        ]
    ))
    A(_make_spacer(D))

    A(_make_para(D,
        "The composition exhibits the classical insurance pattern: the youngest and oldest drivers "
        "exhibit the highest claim frequency, while middle-aged drivers (group 2) have the lowest. "
        "Because group sizes vary by nearly a factor of seven, naive empirical risk minimisation is "
        "dominated by the two large middle groups, motivating the per-group normalisation 1/|I_k⁺| "
        "inside constraint (8)."
    ))

    A(_make_heading(D, "4.2 Models and Baselines", 2))
    A(_make_para(D,
        "We compare two estimators that share the same LP backbone, isolating the effect of the "
        "multi-group fairness constraints."
    ))
    A(_make_para(D,
        "Setting ρ = 0 removes the distributional-robustness penalty ρ‖w‖_1 from constraints "
        "(7), (9), and (10), and sending ζ → ∞ deactivates the fairness constraint (8). The LP "
        "reduces to a class-weighted hinge-loss linear classifier and serves as our "
        "predictive-performance baseline.",
        bold_prefix="Unconstrained Wasserstein SVM (ρ = 0, ζ = ∞). "
    ))
    A(_make_para(D,
        "This is the full multi-group estimator of Theorem 3.5 with Wasserstein radius ρ = 0.1 "
        "and fairness tolerance ζ = 1.5. By Proposition 3.4, H_kl ≥ 1 for every classifier, so "
        "ζ = 1.5 enforces the ratio bound H_kl ≤ 2.5 for all K(K − 1) = 20 ordered pairs "
        "(k, l) ∈ 𝒫_K. The choice ζ = 1.5 is moderate; a systematic sweep over (ρ, ζ) "
        "is left to future work.",
        bold_prefix="HDRFC-K (ρ = 0.1, ζ = 1.5). "
    ))
    A(_make_para(D,
        "Both estimators use the L1 norm on the feature space, which preserves the problem as a "
        "pure linear program; an L2 ball would lift the formulation to a second-order cone program "
        "and is left for future work. Both estimators use inverse-frequency class weights "
        "c_i = N/(2N_+) for y_i = +1 and c_i = N/(2N_−) for y_i = −1 in the objective. This "
        "deviates from the uniform-weight objective stated in Theorem 3.5 but is necessary in "
        "practice because the 5% positive base rate causes the unweighted LP to collapse to the "
        "all-negative classifier w = 0, b < 0. The class weights enter only the objective; "
        "constraints (7)–(10) are unchanged, and all dual and structural results from Section 3 "
        "carry over verbatim."
    ))
    A(_make_para(D,
        "A second implementation choice concerns the dual variables λ^{kl}. A direct reading of "
        "Theorem 3.5 would suggest one variable per training point per pair, yielding "
        "Σ_{(k,l)} N entries. Because (9) and (10) involve only indices in I_k⁺ and I_l⁺, we "
        "instantiate λ^{kl} with dimension |I_k⁺| + |I_l⁺| — this is exact, not an approximation, "
        "and reduces the count of fairness-block variables from approximately 336,000 to "
        "approximately 63,000 on our training fold."
    ))
    A(_make_para(D,
        "The LPs are formulated in CVXPY (Diamond & Boyd, 2016) and solved with CLARABEL "
        "(Goulart & Chen, 2024). Both models reach optimal solver status. End-to-end solve times "
        "are measured on a single workstation and reported with the predictive results; they "
        "should be interpreted as orders of magnitude rather than calibrated benchmarks."
    ))

    A(_make_heading(D, "4.3 Evaluation Protocol and Metrics", 2))
    A(_make_para(D,
        "All metrics are computed on the held-out test fold of 4,000 policies, using the trained "
        "(w, b) to assign ŷ_i = sign(wᵀ x_i + b). We report three families of quantities:"
    ))
    A(_make_para(D,
        "1.  Predictive performance. Plain classification accuracy (1/N_te) Σ_i 1{ŷ_i = y_i} and "
        "balanced accuracy (1/2)(TPR + TNR), the latter being more informative under the 5% "
        "positive base rate."
    ))
    A(_make_para(D,
        "2.  Per-group equal-opportunity statistics. The group-conditional true positive rate "
        "TPR_k = P(ŷ = +1 | y = +1, A = k) for each k ∈ {1,…,5}, the maximum pairwise TPR "
        "gap Δ := max_{k,l} |TPR_k − TPR_l|, and the full K × K pairwise difference matrix "
        "TPR_k − TPR_l which exposes the directional structure of any residual disparity."
    ))
    A(_make_para(D,
        "3.  Computational cost. Wall-clock LP solve time, reported in seconds."
    ))
    A(_make_para(D,
        "The TPR-based metrics correspond directly to the equal-opportunity criterion that "
        "constraints (8)–(10) are designed to enforce; demographic-parity and equalised-odds "
        "variants are outside the scope of this paper."
    ))

    return elems


def build_section_5(doc):
    D = doc
    elems = []
    A = elems.append

    A(_make_heading(D, "5. Results and Discussion", 1))

    A(_make_heading(D, "5.1 Predictive Performance", 2))
    A(_make_para(D,
        "Table 2 summarises the headline comparison. Accuracy decreases by 0.0098 (from 0.5485 "
        "to 0.5387) when fairness and robustness are imposed, while balanced accuracy improves "
        "by 0.0161 (from 0.5753 to 0.5914). The sign reversal between the two metrics is "
        "informative: the unconstrained SVM exploits the 5% class imbalance to reach slightly "
        "higher raw accuracy by under-predicting positives, whereas HDRFC-K — which is forced "
        "by the equal-opportunity constraints to lift TPR in the larger groups — recovers a "
        "more balanced operating point. Solve time grows from 3.1 s to 16.0 s, a factor of "
        "5.2×, consistent with the additional K(K − 1) = 20 blocks of fairness constraints."
    ))

    A(_make_caption(D, "Table 2. Test-set comparison on N_te = 4,000 held-out policies."))
    A(_make_table(D,
        ["Model", "Accuracy", "Balanced acc.", "Max TPR gap Δ", "Solve time"],
        [
            ["Unconstrained SVM (ρ = 0, ζ = ∞)", "0.5485", "0.5753", "0.3073", "3.1 s"],
            ["HDRFC-K (ρ = 0.1, ζ = 1.5)",           "0.5387", "0.5914", "0.1232", "16.0 s"],
            ["Change",                                          "−0.0098", "+0.0161", "−0.1842 (−59.9%)", "—"],
        ]
    ))
    A(_make_spacer(D))

    A(_make_para(D,
        "In absolute terms both models leave substantial predictive headroom on this dataset; "
        "freMTPL2freq is well known to be a noisy, high-class-imbalance benchmark for which "
        "linear classifiers without exposure modelling routinely report balanced accuracies in "
        "the high fifties. We therefore interpret the 1.6 percentage-point gain in balanced "
        "accuracy as a side-effect of the class-aware constraint geometry rather than as "
        "evidence of a fundamentally stronger classifier. The plain-accuracy decrease of roughly "
        "one percentage point is the ‘fairness tax’ in this experiment."
    ))

    A(_make_heading(D, "5.2 Fairness Results", 2))
    A(_make_para(D,
        "The maximum pairwise TPR gap drops from Δ = 0.3073 under the unconstrained SVM to "
        "Δ = 0.1232 under HDRFC-K, a reduction of 59.9%. Table 3 reports the full per-group "
        "TPR profile."
    ))

    A(_make_caption(D, "Table 3. Per-group true positive rate on the test fold."))
    A(_make_table(D,
        ["k", "Age", "TPR_k (SVM)", "TPR_k (HDRFC-K)", "Change"],
        [
            ["1", "18–25", "0.8667", "0.7333", "−0.1333"],
            ["2", "26–40", "0.5593", "0.6102", "+0.0508"],
            ["3", "41–55", "0.5921", "0.6711", "+0.0789"],
            ["4", "56–70", "0.6000", "0.6286", "+0.0286"],
            ["5", "71+",        "0.6000", "0.6667", "+0.0667"],
        ]
    ))
    A(_make_spacer(D))

    A(_make_para(D,
        "The unconstrained SVM exhibits a pronounced asymmetry: the youngest group (18–25) "
        "has TPR_1 = 0.8667 while the four older groups lie in the narrow band [0.5593, 0.6000]. "
        "The fairness constraints (8) act in the direction one would predict from Theorem 3.5: "
        "TPR_1 is reduced by 0.1333 towards the bulk of the distribution, while "
        "TPR_2, …, TPR_5 each rise by between 0.0286 and 0.0789. After projection, all five "
        "group TPRs lie in [0.6102, 0.7333], so no group is left below 0.61."
    ))
    A(_make_para(D,
        "The pairwise structure makes the geometry concrete. Under the unconstrained SVM, every "
        "pair involving group 1 shows a large positive gap: TPR_1 − TPR_2 = +0.307, "
        "TPR_1 − TPR_3 = +0.275, TPR_1 − TPR_4 = +0.267, TPR_1 − TPR_5 = +0.267, while all "
        "C(4,2) = 6 pairs among groups {2, 3, 4, 5} lie within ±0.041. The disparity is "
        "therefore essentially a one-versus-rest pattern driven by the youngest drivers. "
        "Under HDRFC-K, the analogous pairwise differences contract: "
        "TPR_1 − TPR_2 = +0.123, TPR_1 − TPR_4 = +0.105, TPR_1 − TPR_5 = +0.067, "
        "TPR_1 − TPR_3 = +0.062, TPR_3 − TPR_2 = +0.061, with all remaining pairs among "
        "groups {2, 3, 4, 5} within ±0.057. The new maximum, 0.123, again involves group 1 "
        "but is now less than half of the largest baseline gap (0.307). None of the realised "
        "pairwise ratios approach the slack bound H_kl ≤ 2.5 implied by ζ = 1.5; the "
        "constraints in (8) are therefore not tight at the optimum, suggesting that tighter "
        "ζ values would still be feasible at this ρ."
    ))

    A(_make_heading(D, "5.3 Discussion", 2))
    A(_make_para(D,
        "The empirical trade-off on this benchmark is mild: a 59.9% reduction in maximum TPR "
        "gap is purchased at the cost of 0.98 percentage points of plain accuracy, while "
        "balanced accuracy actually improves. This is consistent with the regime described by "
        "the LP geometry of Theorem 3.5 when the unconstrained classifier is already near a "
        "fair operating point along most pairs and the disparity is concentrated in a small "
        "number of groups (here, group 1). In such cases the projection induced by constraint "
        "(8) may act on a low-dimensional slice of the decision boundary and the predictive "
        "surplus elsewhere could be largely preserved. We caution that this favourable behaviour "
        "is not guaranteed by the theory; on datasets where the unconstrained optimum is far "
        "from any equal-opportunity solution, the price of fairness can be substantially larger.",
        bold_prefix="Accuracy–fairness trade-off. "
    ))
    A(_make_para(D,
        "The two regularisers play complementary roles. The Wasserstein radius ρ enters "
        "constraints (7), (9), and (10) through the ρ‖w‖_1 term and acts as an L1 penalty "
        "that hedges against distributional shift; the fairness tolerance ζ enters only "
        "constraint (8) and controls how close the solution is forced to perfect "
        "equal-opportunity (ζ = 0). At our chosen operating point (ρ, ζ) = (0.1, 1.5) the "
        "fairness constraints are slack at the optimum, suggesting that the force shaping the "
        "per-group TPR profile may reflect the interaction between ρ and ζ rather than (8) "
        "alone. A systematic Pareto sweep over (ρ, ζ), including the limit ζ → 0, is the "
        "natural next experimental step and would let us trace the full accuracy–fairness "
        "frontier predicted by the LP.",
        bold_prefix="Role of ρ and ζ. "
    ))
    A(_make_para(D,
        "Several caveats apply. First, the results are obtained on a single dataset and a "
        "single train/test split; we report point estimates rather than confidence intervals, "
        "and any claim about the relative ranking of the two estimators should therefore be "
        "regarded as preliminary until cross-validated on multiple folds and replicated on "
        "additional benchmarks. Second, the inverse-frequency class weighting, while a standard "
        "remedy for class imbalance, departs from the uniform-weight objective in Theorem 3.5; "
        "although the constraint structure is unchanged, the precise duality argument underlying "
        "the LP would need to be re-stated for the weighted case to be fully rigorous. Third, "
        "the L1 norm is a modelling choice driven by computational tractability — an L2 (SOCP) "
        "variant would be preferable when the feature scaling is heterogeneous, and we have not "
        "characterised how much of the observed fairness gain is specific to L1 geometry. "
        "Fourth, the sensitive attribute is excluded from x_i but its proxies (region, area "
        "code, bonus-malus) remain, so the residual TPR gap of 0.123 in group 1 may reflect "
        "proxy leakage that the linear-LP framework cannot, by construction, fully eliminate. "
        "Fifth, our solve-time figures are wall-clock measurements on a single machine and "
        "should not be read as a benchmark of CLARABEL or of LP-based fair learning more "
        "broadly. Finally, the experiment uses K = 5 groups; the LP scales as K(K − 1) in "
        "the number of fairness blocks, and the empirical scaling behaviour for larger K "
        "remains to be characterised.",
        bold_prefix="Limitations. "
    ))

    return elems


# ── new APA references to append ──────────────────────────────────────────

NEW_REFS = [
    ("Diamond, J., & Boyd, S. (2016). CVXPY: A Python-embedded modeling language for "
     "convex optimization. Journal of Machine Learning Research, 17(83), 1–5. "
     "https://www.jmlr.org/papers/v17/15-408.html"),
    ("Duval, C., & Charpentier, A. (2020). freMTPL2freq [Dataset]. OpenML. "
     "https://www.openml.org/d/41214"),
    ("Goulart, P., & Chen, Y. (2024). CLARABEL: An interior-point solver for conic "
     "optimization problems (Version 0.9) [Software]. "
     "https://github.com/oxfordcontrol/Clarabel.rs"),
]


# ── main ──────────────────────────────────────────────────────────────────

def main():
    doc = Document("Combined_Paper.docx")
    body = doc.element.body

    # 1. Locate key paragraphs
    def find_idx(text):
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == text:
                return i
        return None

    sec4_idx = find_idx("4. Dataset and Experimental Setup")
    sec6_idx = find_idx("6. Conclusion")
    assert sec4_idx is not None and sec6_idx is not None

    print(f"Section 4 at [{sec4_idx}], Section 6 at [{sec6_idx}]")

    # 2. Remove old sections 4 and 5 (indices sec4_idx .. sec6_idx-1)
    to_remove = [p._element for p in doc.paragraphs[sec4_idx:sec6_idx]]
    anchor = doc.paragraphs[sec6_idx]._element   # element to insert before

    for el in to_remove:
        el.getparent().remove(el)

    # 3. Build new content elements (all detached)
    new_elems = build_section_4(doc) + build_section_5(doc)

    # 4. Insert new elements before Section 6 heading
    # anchor may have moved if it was part of removed elems — re-fetch
    # (sec6 was NOT in to_remove since we sliced [sec4:sec6])
    for el in new_elems:
        anchor.addprevious(el)

    # 5. Append new references at end of document
    last_para = doc.paragraphs[-1]
    last_el = last_para._element
    for ref_text in NEW_REFS:
        p = doc.add_paragraph(ref_text, style="Normal")
        new_el = p._element
        body.remove(new_el)
        last_el.addnext(new_el)
        last_el = new_el

    # 6. Save
    out = "papers/Combined_Paper_Revised.docx"
    doc.save(out)
    print(f"Saved: {out}")

    # 7. Quick verification
    doc2 = Document(out)
    headings = [(i, p.text) for i, p in enumerate(doc2.paragraphs)
                if p.style.name.startswith("Heading") and p.text.strip()]
    print("\nHeading structure:")
    for idx, h in headings:
        print(f"  [{idx:03d}] {h}")


if __name__ == "__main__":
    main()
