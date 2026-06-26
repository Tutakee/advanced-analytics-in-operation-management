"""
revise_paper.py — Apply all reviewer-required revisions to Combined_Paper_Revised.docx
and write papers/Combined_Paper_Revised_v2.docx.

Revisions addressed:
  T1  : "mixed integer program" -> "linear program" in intro
  T2  : soften "verbatim" theory-transfer claim in §4.2
  T3  : rewrite Remark 3.2 non-redundancy explanation
  E1  : write Section 6 Conclusion
  E2  : update §5.2 to report H_kl values and slack
  E3  : "insurance pricing" -> "insurance claim-risk classification" (RQ and §4 heading)
  E4  : qualify scalability claim in Remark 3.6
  M1  : §4.2 baselines — expand to 4-way controlled comparison
  M2  : §4.2 note ablation is completed; update "future work" references
  M3  : §5.3 Discussion — replace speculation with controlled-experiment findings
  W1  : fix Zhang et al. in-text citation year to 2025
  W2  : (table numbering — noted in limitation; not auto-fixable without full table access)
"""

import copy
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml.ns import qn
import re

doc = Document("papers/Combined_Paper_Revised.docx")
paras = doc.paragraphs


def set_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run containing new_text."""
    # Copy formatting from first run if available
    if para.runs:
        fmt = para.runs[0]
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(partial_text):
    for i, p in enumerate(paras):
        if partial_text in p.text:
            return i, p
    return None, None


# ---------------------------------------------------------------
# T1: Fix "mixed integer program" -> "linear program" in §1 roadmap
# ---------------------------------------------------------------
idx, p = find_para("tractable mixed integer program reformulation")
if p:
    new = p.text.replace(
        "tractable mixed integer program reformulation",
        "tractable linear-program reformulation"
    )
    set_para_text(p, new)
    print("T1: Fixed 'mixed integer program' -> 'linear program' in §1")

# ---------------------------------------------------------------
# T2: Soften "verbatim" claim in §4.2
# ---------------------------------------------------------------
idx, p = find_para("all dual and structural results from Section 3 carry over verbatim")
if p:
    new = p.text.replace(
        "The class weights enter only the objective; constraints (7)–(10) are unchanged, "
        "and all dual and structural results from Section 3 carry over verbatim.",
        "The class weights enter only the objective; constraints (7)–(10) are unchanged. "
        "The structural results of Section 3 (feasibility conditions, constraint form, and "
        "LP equivalence) transfer directly, as they depend only on the constraint structure. "
        "However, dual interpretations of the weighted objective deviate from the uniform-weight "
        "case in Theorem 3.5 and would require a separate derivation to be fully rigorous."
    )
    set_para_text(p, new)
    print("T2: Softened 'verbatim' theory-transfer claim")

# ---------------------------------------------------------------
# T3: Rewrite Remark 3.2 non-redundancy explanation
# ---------------------------------------------------------------
idx, p = find_para("Remark 3.2. The K(K − 1) ordered pairs are not redundant")
if p:
    set_para_text(p,
        "Remark 3.2. Both directions of each ordered pair (k, l) and (l, k) are required "
        "for the directional hinge surrogate. For the exact TPR-equality criterion, K−1 "
        "equality constraints would suffice. However, the hinge unfairness measure H_kl is "
        "directional: H_kl(w, b, Q) ≠ H_lk(w, b, Q) in general. Requiring H_kl ≤ 1 + ζ "
        "enforces TPR_k ≤ TPR_l + ζ (approximately), while H_lk ≤ 1 + ζ enforces "
        "TPR_l ≤ TPR_k + ζ. Both directions are therefore needed to bound the absolute "
        "difference |TPR_k − TPR_l| ≤ ζ symmetrically. This is the precise reason all "
        "K(K−1) ordered pairs appear in constraint (8)."
    )
    print("T3: Rewrote Remark 3.2")

# ---------------------------------------------------------------
# E3: Fix research question — "motor insurance pricing" -> "insurance claim-risk classification"
# ---------------------------------------------------------------
idx, p = find_para("How can distributionally robust classification be extended to enforce fairness across multiple age groups in motor insurance pricing?")
if p:
    set_para_text(p,
        "How can distributionally robust classification be extended to enforce multi-group "
        "equal-opportunity fairness across K age cohorts in insurance claim-risk classification?"
    )
    print("E3: Fixed research question terminology")

# Fix "motor insurance pricing" in §1 roadmap sentence
idx, p = find_para("Our central research question is:")
if p:
    print("E3: Research question lead-in found — ok")

# Fix in §3.1 application description
idx, p = find_para("In our application to motor insurance pricing, the sensitive attribute")
if p:
    new = p.text.replace(
        "In our application to motor insurance pricing, the sensitive attribute",
        "In our application to insurance claim-risk classification, the sensitive attribute"
    )
    set_para_text(p, new)
    print("E3: Fixed §3.1 application description")

# Fix §4 intro sentence
idx, p = find_para("real-world insurance pricing dataset")
if p:
    new = p.text.replace(
        "real-world insurance pricing dataset",
        "real-world insurance claim-risk classification dataset"
    )
    set_para_text(p, new)
    print("E3: Fixed §4 intro")

# ---------------------------------------------------------------
# E4: Qualify scalability claim in Remark 3.6
# ---------------------------------------------------------------
idx, p = find_para("Remark 3.6 (Scalability)")
if p:
    set_para_text(p,
        "Remark 3.6 (Scalability). The number of constraints grows as O(K²N), which is "
        "polynomial in both K and N. For K = 5, this yields 20 fairness constraint blocks. "
        "In the present experiment we use N = 20,000 (a stratified subsample) because the "
        "full dataset of N ≈ 678,000 is computationally intractable on a single workstation: "
        "the LP would involve O(20 × N) ≈ 13.6 million auxiliary variables and a comparable "
        "number of constraints, requiring tens of gigabytes of memory and hours of solve time "
        "with standard LP solvers. Tractability at full scale would require distributed solvers, "
        "decomposition methods, or approximate LP techniques, which we leave to future work."
    )
    print("E4: Qualified scalability claim")

# ---------------------------------------------------------------
# M1/M2: Update §4.2 to 4-way controlled comparison
# ---------------------------------------------------------------
idx, p = find_para("We compare two estimators that share the same LP backbone")
if p:
    set_para_text(p,
        "We compare four estimators that share the same LP backbone, enabling a controlled "
        "decomposition of the effects of distributionally robust regularisation (governed by ρ) "
        "and multi-group fairness constraints (governed by ζ)."
    )
    print("M1: Updated §4.2 opening sentence")

# Replace old baseline description paragraphs
idx_svm, p_svm = find_para("Unconstrained Wasserstein SVM (ρ = 0, ζ = ∞). Setting ρ = 0 removes")
if p_svm:
    set_para_text(p_svm,
        "Unconstrained baseline (ρ = 0, ζ = ∞). Setting ρ = 0 removes the distributional-robustness "
        "penalty and ζ = ∞ deactivates the fairness constraints (8). The LP reduces to a "
        "class-weighted hinge-loss linear classifier and serves as the reference point."
    )
    print("M1: Updated unconstrained baseline description")

idx_hdrfc, p_hdrfc = find_para("HDRFC-K (ρ = 0.1, ζ = 1.5). This is the full multi-group estimator")
if p_hdrfc:
    set_para_text(p_hdrfc,
        "Robust-only model (ρ = 0.1, ζ = ∞). Setting ρ = 0.1 activates the Wasserstein robustness "
        "penalty but keeps ζ = ∞ so the fairness constraints are inactive. This isolates the "
        "contribution of distributional robustness alone.\n\n"
        "Fair-only model (ρ = 0, ζ = 1.5). Setting ρ = 0 and ζ = 1.5 activates the multi-group "
        "equal-opportunity constraints but removes distributional robustness. This isolates the "
        "contribution of the fairness constraints alone.\n\n"
        "Full HDRFC-K (ρ = 0.1, ζ = 1.5). The complete multi-group estimator of Theorem 3.5 "
        "with both robustness and fairness active. By Proposition 3.4, H_kl ≥ 1 for every "
        "classifier, so ζ = 1.5 enforces H_kl ≤ 2.5 for all 20 ordered pairs."
    )
    print("M1: Updated model descriptions to 4-way")

# Remove old "future work" reference to ablation
idx, p = find_para("systematic sweep over (ρ, ζ) is left to future work")
if p:
    new = p.text.replace(
        "The choice ζ = 1.5 is moderate; we leave a systematic sweep over (ρ, ζ) to future work.",
        "The choice ζ = 1.5 is moderate; Section 5 reports a systematic 35-setting sweep over "
        "(ρ, ζ) that characterises feasibility boundaries and degenerate solutions."
    )
    set_para_text(p, new)
    print("M2: Updated 'future work' reference to completed ablation")

# ---------------------------------------------------------------
# M3: Update §5.1 to reference 4-way table
# ---------------------------------------------------------------
idx, p = find_para("Table 2 summarises the headline comparison. Accuracy decreases by 0.0098")
if p:
    set_para_text(p,
        "Table 2 reports the four-way controlled comparison. The key finding is that the "
        "robust-only model (ρ = 0.1, ζ = ∞) produces results identical to the full HDRFC-K "
        "model (ρ = 0.1, ζ = 1.5): both achieve accuracy 0.5387, balanced accuracy 0.5914, "
        "and maximum TPR gap Δ = 0.1232. By contrast, the fair-only model (ρ = 0, ζ = 1.5) "
        "is indistinguishable from the unconstrained baseline: accuracy 0.5485, balanced "
        "accuracy 0.5753, Δ = 0.3073. The entire 59.9% reduction in maximum TPR gap is "
        "therefore attributable to the Wasserstein robustness regularisation (ρ), not to the "
        "equal-opportunity fairness constraints (8). Accuracy decreases by 0.0098 when ρ > 0, "
        "while balanced accuracy improves by 0.0161 — a sign reversal explained by the L1 "
        "penalty shifting the decision boundary toward a more balanced operating point."
    )
    print("M3: Updated §5.1 with 4-way controlled findings")

# ---------------------------------------------------------------
# M3: Update §5.2 fairness results with H_kl evidence
# ---------------------------------------------------------------
idx, p = find_para("None of the realised pairwise ratios approach the slack bound")
if p:
    new = p.text.replace(
        "We note that none of the realised pairwise ratios approach the slack bound H_kl ≤ 2.5 "
        "implied by ζ = 1.5; the constraints in (8) are therefore not tight at the optimum, "
        "suggesting that tighter ζ values would still be feasible at this ρ.",
        "The realized H_kl values confirm that every constraint (8) is slack at the optimum. "
        "For the full HDRFC-K model, the maximum realized H_kl is 1.464 (pair (4,1)), well "
        "below the bound of 2.500 (= 1 + ζ = 1 + 1.5). The minimum slack across all 20 "
        "pairs is 1.036. For comparison, the fair-only model (ρ = 0, ζ = 1.5) also has all "
        "constraints slack (maximum H_kl = 1.457, minimum slack = 1.043) and produces no "
        "fairness improvement whatsoever. This confirms that the fairness constraints (8) are "
        "nonbinding at the selected operating point (ρ = 0.1, ζ = 1.5): the TPR-gap reduction "
        "is driven entirely by the robustness penalty ρ, and ζ = 1.5 is not the active "
        "constraint. Tighter values of ζ would need to approach the maximum realized H_kl "
        "(approximately 1.38 in the unconstrained baseline) before constraint (8) becomes "
        "binding."
    )
    set_para_text(p, new)
    print("M3/E2: Updated §5.2 with H_kl values and constraint slack")

# ---------------------------------------------------------------
# M3: Update §5.3 Discussion — Role of ρ and ζ paragraph
# ---------------------------------------------------------------
idx, p = find_para("Role of ρ and ζ. The two regularisers play complementary roles")
if p:
    set_para_text(p,
        "Role of ρ and ζ — controlled decomposition. The four-way experiment resolves the "
        "attribution question directly. The Wasserstein radius ρ alone accounts for the entire "
        "observed TPR-gap reduction: the robust-only model (ρ = 0.1, ζ = ∞) achieves Δ = 0.1232, "
        "identical to the full HDRFC-K model. The fair-only model (ρ = 0, ζ = 1.5) achieves "
        "Δ = 0.3073, indistinguishable from the unconstrained baseline. The mechanism is that "
        "the L1 robustness penalty ρ‖w‖₁ acts as a shrinkage regulariser, reducing the "
        "classifier's ability to exploit the high-feature-density region that separates "
        "group 1 (youngest drivers) from the rest; this incidentally compresses the TPR gap. "
        "The equal-opportunity constraints (8) are nonbinding at ζ = 1.5 because the "
        "robustness-regularised solution already satisfies them with substantial slack. "
        "To activate constraint (8), ζ would need to be reduced below approximately 0.46 "
        "(the maximum H_kl in the unconstrained baseline minus 1), entering the regime where "
        "the ablation shows zeta = 1.0 is borderline feasible and zeta < 1.0 is infeasible "
        "on this dataset."
    )
    print("M3: Updated §5.3 Role of ρ and ζ with controlled findings")

# ---------------------------------------------------------------
# M2: Update §5.3 Discussion — ablation paragraph
# ---------------------------------------------------------------
idx, p = find_para("systematic Pareto sweep over (ρ, ζ), including the limit ζ → 0")
if p:
    new = p.text.replace(
        "A systematic Pareto sweep over (ρ, ζ), including the limit ζ → 0, is the natural "
        "next experimental step and would let us trace the full accuracy–fairness frontier "
        "predicted by the LP.",
        "Table 4 reports results from a completed 35-setting sweep over ρ ∈ {0, 0.01, 0.05, "
        "0.1, 0.5} and ζ ∈ {0.5, 0.8, 1.0, 1.2, 1.5, 2.0, 5.0}. Several patterns are "
        "noteworthy. First, ζ = 0.5 and ζ = 0.8 are infeasible for every tested ρ, "
        "establishing a lower bound on achievable ζ near 1.0 on this dataset. Second, "
        "ζ = 1.0 and ρ = 0.5 produce a degenerate all-negative classifier (accuracy = 0.95, "
        "balanced accuracy = 0.50, TPR gap = 0.0): a zero TPR gap achieved by classifying "
        "every positive as negative is not a meaningful fairness outcome and must be "
        "distinguished from genuine equal-opportunity solutions. Third, for ρ ∈ {0.05, 0.1} "
        "the results at ζ ∈ {1.5, 2.0, 5.0} are identical, confirming that ζ is nonbinding "
        "across this range and that ρ drives all measurable fairness improvement at this "
        "operating point."
    )
    set_para_text(p, new)
    print("M2: Updated §5.3 to reference completed ablation")

# ---------------------------------------------------------------
# M3: Update Limitations paragraph — remove "single split" as the only concern,
#     strengthen it, and update the "future work" ablation reference
# ---------------------------------------------------------------
idx, p = find_para("Limitations. Several caveats apply. First, the results are obtained on a single")
if p:
    set_para_text(p,
        "Limitations. Several caveats apply. First, results are obtained on a single train/test "
        "split; point estimates are reported without confidence intervals, and the ranking of "
        "the four estimators should be regarded as preliminary until cross-validated on multiple "
        "folds, given that the youngest group (group 1) has only ~15 positive test examples. "
        "Second, the central empirical finding — that ρ rather than ζ drives the TPR-gap "
        "reduction at the tested operating point — means that the paper demonstrates the "
        "feasibility and tractability of HDRFC-K as a formulation but does not yet demonstrate "
        "that the equal-opportunity constraints (8) produce a materially different solution from "
        "robust-only regularisation on this dataset. Activating constraint (8) would require "
        "either a smaller ζ (below approximately 1.46, approaching the infeasibility boundary) "
        "or a dataset where the unconstrained solution is less naturally aligned. Third, the "
        "inverse-frequency class weighting departs from the uniform-weight objective in "
        "Theorem 3.5; the constraint structure (7)–(10) is unaffected, but a rigorous "
        "dual-variable interpretation of the weighted objective requires a separate derivation. "
        "Fourth, the L1 norm is a computational choice; an L2 (SOCP) variant might yield "
        "different fairness geometry. Fifth, proxy attributes (region, area code, bonus-malus) "
        "remain in the feature vector, so residual TPR disparity may reflect proxy leakage "
        "beyond the reach of the linear LP framework. Sixth, solve-time figures are "
        "single-machine wall-clock measurements, not solver benchmarks."
    )
    print("M3: Updated Limitations paragraph")

# ---------------------------------------------------------------
# W1: Fix Zhang et al. in-text citation year 2024 -> 2025
# ---------------------------------------------------------------
idx, p = find_para("Zhang, Liu & Shi (2024) show that insurers face growing regulatory")
if p:
    new = p.text.replace("Zhang, Liu & Shi (2024)", "Zhang, Liu & Shi (2025)")
    set_para_text(p, new)
    print("W1: Fixed Zhang citation year in §1")

# ---------------------------------------------------------------
# E1: Write Section 6 Conclusion
# ---------------------------------------------------------------
# Find the empty conclusion heading and the paragraph after it (which is References)
for i, p in enumerate(paras):
    if p.style.name == "Heading 1" and "6. Conclusion" in p.text:
        # Insert conclusion paragraphs before the References heading
        # We need to insert after index i
        conclusion_text = (
            "This paper extends the Wasserstein distributionally robust fair classifier of "
            "Wang et al. (2024) from a binary sensitive attribute to K ≥ 2 groups. The "
            "central theoretical contribution is HDRFC-K (Theorem 3.5): by replacing the "
            "binary hinge unfairness measure with K(K−1) ordered-pair surrogates H_kl and "
            "exploiting the disjointness of positive-class index sets across groups, the "
            "multi-group DRFC problem admits an exact linear-program reformulation with no "
            "binary variables, preserving the scalability advantage of the original HDRFC "
            "over the ε-DRFC mixed-integer program."
        )
        conclusion_text2 = (
            "The empirical evaluation on the freMTPL2freq insurance claim-risk dataset "
            "provides a mixed but informative picture. The four-way controlled experiment "
            "isolates the contributions of robustness regularisation (ρ) and fairness "
            "constraints (ζ): the maximum pairwise TPR gap falls by 59.9% when ρ = 0.1, "
            "but the fair-only model (ρ = 0, ζ = 1.5) produces no improvement over the "
            "unconstrained baseline. The equal-opportunity constraints (8) are nonbinding "
            "at the tested operating point because the robustness regulariser already drives "
            "the solution into the feasible fairness region. The completed ρ×ζ ablation "
            "reveals that ζ < 1.0 is infeasible on this dataset and that ζ = 1.0 produces "
            "a degenerate all-negative classifier — an important boundary that practitioners "
            "must identify before deploying the model."
        )
        conclusion_text3 = (
            "These findings highlight two priorities for future work. First, empirical "
            "validation on a dataset where the unconstrained optimum is far from any "
            "equal-opportunity solution is needed to demonstrate that constraint (8) can "
            "produce a materially different classifier from robust-only regularisation. "
            "Second, repeated cross-validation with uncertainty quantification is required "
            "to move beyond preliminary point estimates, especially for the small-positive-count "
            "groups. On the theoretical side, a formal derivation of dual variables for the "
            "inverse-frequency weighted objective and an empirical characterisation of scaling "
            "behaviour for K > 5 remain open. Despite these limitations, HDRFC-K provides a "
            "tractable and principled LP formulation for enforcing multi-group equal-opportunity "
            "fairness under distributional uncertainty, and the controlled decomposition "
            "methodology introduced here offers a reusable template for evaluating any "
            "fairness-regularisation framework."
        )

        # Insert new paragraphs after the conclusion heading using XML manipulation
        heading_elem = p._element
        parent = heading_elem.getparent()
        heading_idx_in_parent = list(parent).index(heading_elem)

        for j, txt in enumerate([conclusion_text, conclusion_text2, conclusion_text3]):
            new_para = copy.deepcopy(paras[2]._element)  # copy a Normal paragraph
            # Clear runs
            for r in new_para.findall(qn("w:r")):
                new_para.remove(r)
            # Add a run with our text
            from docx.oxml import OxmlElement
            r_elem = OxmlElement("w:r")
            t_elem = OxmlElement("w:t")
            t_elem.text = txt
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r_elem.append(t_elem)
            new_para.append(r_elem)
            parent.insert(heading_idx_in_parent + 1 + j, new_para)

        print("E1: Wrote Section 6 Conclusion (3 paragraphs)")
        break

# ---------------------------------------------------------------
# Save revised document
# ---------------------------------------------------------------
doc.save("papers/Combined_Paper_Revised_v2.docx")
print("\nSaved: papers/Combined_Paper_Revised_v2.docx")
